from selenium import webdriver
import pyautogui
import docx
import pytesseract
import tkinter as tk
from tkinter import *
from PIL import ImageTk, Image, ImageDraw
import pygame
import time
import os
from itertools import count
def Hwit_vit():

    rootH1 = Toplevel()
    rootH1.state('zoomed')
    rootH1.title('वर्चुअल पुलिस स्टेशन1')
    rootH1.configure(bg='#ffeb99')
    FaH1 = tk.Frame(rootH1, borderwidth=40, bg="black", relief=tk.SUNKEN)
    lbH1 = ImageLabel(rootH1)
    lbH1.pack(side=tk.TOP)
    lbH1.load('Y_giff.gif')
    pygame.init()
    pygame.mixer.music.load('h2.mp3')
    pygame.mixer.music.play()

    FaH1.pack()
    FaH2 = tk.Frame(rootH1, borderwidth=40, bg="#ffeb99", relief=tk.SUNKEN)
    lbH2 = tk.Label(rootH1, text="आप अपने बारे में बताये", bg="#ffeb99")
    lbH2.pack(side=tk.TOP)

    butH_witness= tk.Button(FaH2, text="सूचक", width='63', height='5',command=Hper_wit)
    butH_witness.pack(side=tk.LEFT, padx=150, fill=tk.X)
    butH_victim= tk.Button(FaH2, text="पीड़ित", width='63', height='5',command=Hper_wit)
    butH_victim.pack(side=tk.RIGHT, padx=20, fill=tk.X)
    FaH2.pack(side=tk.BOTTOM, fill=tk.BOTH)
    rootH1.mainloop()

def Hper_wit():
    rootH4 = Toplevel()
    rootH4.state('zoomed')
    rootH4.title('वर्चुअल पुलिस स्टेशन2')
    rootH4.configure(bg='#ffeb99')


    labelH3_0 = tk.Label(rootH4, text="अपने व्यक्तिगत विवरण प्रदान करें", width=40,bg='#ffeb99' ,font=("bold", 23)).grid(row=0, column=0)
    lb31H3 = ImageLabel(rootH4)
    lb31H3.grid(row=0, column=4)
    lb31H3.load('Y_giff.gif')


    labelH3_1 = tk.Label(rootH4, text="अपना नाम बताएं।", width=40, bg='#ffeb99', font=("bold", 13),padx=10,pady=15).grid(row=1, column=0)

    entryH3_1 = tk.Entry(rootH4)
    entryH3_1.grid(row=1, column=1)

    def Hplay1_wt():
        pygame.mixer.music.load("a5.mp3")
        pygame.mixer_music.play()
    butH3_q1 = tk.Button(rootH4, text="सुने", width='10', height='1',command=Hplay1_wt).grid(row=1, column=2)
    butH3_r1 = tk.Button(rootH4, text="बोलें", width='10', height='1').grid(row=1, column=3)


    pygame.init()
    pygame.mixer.music.load('h1.mp3')
    pygame.mixer.music.play()
    labelH3_2 = tk.Label(rootH4, text="अपना फोन नंबर बतायें", bg='#ffeb99', width=40, font=("bold", 13),padx=10,pady=15).grid(row=2,column=0)


    entryH3_2 = tk.Entry(rootH4)
    entryH3_2.grid(row=2, column=1)
    def Hplay2_wt():
        pygame.mixer.music.load('a10.mp3')
        pygame.mixer_music.play()
    butH3_q2 = tk.Button(rootH4, text="सुने", width='10', height='1',command=Hplay2_wt).grid(row=2, column=2)
    butH3_r2 = tk.Button(rootH4, text="बोलें", width='10', height='1').grid(row=2,column=3)

    labelH3_3 = tk.Label(rootH4, text="अपना लिंग बतायें", bg='#ffeb99', width=20, font=("bold", 13),padx=10,pady=15).grid(row=3, column=0)
    yH3 = tk.IntVar()
    tk.Radiobutton(rootH4, text="पुरुष", padx=5, variable=yH3, value=1, anchor=W, bg='#ffeb99', font=("bold", 12)).grid(
        row=3, column=1)
    tk.Radiobutton(rootH4, text="स्री", padx=5, variable=yH3, value=2, anchor=W, bg='#ffeb99', font=("bold", 12)).grid(
        row=3, column=2)
    tk.Radiobutton(rootH4, text="अन्य", padx=5, variable=yH3, value=3, anchor=W, bg='#ffeb99', font=("bold", 12)).grid(
        row=3, column=3)
    labelH3_4 = tk.Label(rootH4, text="अपना पता बताएं", bg='#ffeb99', width=20, font=("bold", 13),padx=10,pady=15).grid(row=4, column=0)

    entryH3_4 = tk.Entry(rootH4)
    entryH3_4.grid(row=4, column=1)
    def Hplay4_wt():
        pygame.mixer.music.load('a9.mp3')
        pygame.mixer_music.play()
    butH3_q4 = tk.Button(rootH4, text="सुने", width='10', height='1',command=Hplay4_wt).grid(row=4, column=2)
    butH3_r4 = tk.Button(rootH4, text="बोलें", width='10', height='1').grid(row=4, column=3)

    labelH3_5 = tk.Label(rootH4, text="घटना का विवरण प्रदान करें", bg='#ffeb99', width=30, font=("bold", 13),padx=10,pady=15).grid(row=5, column=0)

    entryH3_5 = tk.Entry(rootH4)
    entryH3_4.grid(row=5, column=1)
    def Hplay5_wt():
        pygame.mixer.music.load('a14.mp3')
        pygame.mixer_music.play()
    butH3_q5 = tk.Button(rootH4, text="सुने", width='10', height='1',command=Hplay5_wt).grid(row=5, column=2)
    butH3_r5 = tk.Button(rootH4, text="बोलें", width='10', height='1').grid(row=5, column=3)

    labelH3_6 = tk.Label(rootH4, text="अपराधी का विवरण करें", bg='#ffeb99', width=25, font=("bold", 13),padx=10,pady=15).grid(row=6, column=0)

    entryH3_6 = tk.Entry(rootH4)
    entryH3_6.grid(row=6, column=1)
    def play6_wt():
        pygame.mixer.music.load("a20.mp3")
        pygame.mixer_music.play()

    butH3_q6 = tk.Button(rootH4, text="सुने", width='10', height='1',command=play6_wt).grid(row=6, column=2)
    butH3_r6 = tk.Button(rootH4, text="बोलें", width='10', height='1').grid(row=6, column=3)

    labelH3_7 = tk.Label(rootH4, text="जन्म तिथी", bg='#ffeb99', width=20, font=("bold", 13),padx=10,pady=15).grid(row=7, column=0)

    entryH3_7 = tk.Entry(rootH4)
    entryH3_7 .grid(row=7, column=1)


    butH3_q7 = tk.Button(rootH4, text="सुने", width='10', height='1').grid(row=7, column=2)
    butH3_r7 = tk.Button(rootH4, text="बोलें", width='10', height='1').grid(row=7, column=3)

    # label_8 = tk.Label(root2, text=" Adhar Number", bg='#00cc99', width=8, font=("bold", 13)).grid(row=8, column=0)
    # x12 = StringVar()
    # entry_8 = tk.Entry(root2, textvariable=x12).grid(row=8, column=1)
    # but_q8 = tk.Button(root2, text="question8", width='10', height='1').grid(row=8, column=2)
    # but_r8 = tk.Button(root2, text="record8", width='10', height='1').grid(row=8, column=3)

    butH3_submit = tk.Button(rootH4, text="पीछे जाये", width='12', height='3', padx=5).grid(row=8, column=0)
    # but3_submit = tk.Button(root4, text="Add witness", width='12', height='3', padx=5).grid(row=8, column=1)

    def Hsub_m():
        Hy1 = entryH3_1.get()

        Hy2 = entryH3_2.get()

        Hy4 = entryH3_4.get()

        Hy5 = entryH3_5.get()

        Hy6 = entryH3_6.get()

        Hy7 = entryH3_7.get()



    def Hu():
        Hsub_m()
        adhar()

    butH3_submit = tk.Button(rootH4, text="आगे जाये", width='12', height='3', padx=5, command=Hu).grid(row=8, column=1)


    rootH4.mainloop()

from pyttsx3 import engine
# x5 =''
# x6 =''
# x7 =''
# x8 =''
# x9 =''
# x10 =''
# x11 =''
# x12 =''
# x15 =''
# x16 =''
# x17 =''
# x18 =''
# x19 =''
# x20 =''
# x22 =''
# y1 =''
# y2 =''
# y3 =''
# y4 =''

class ImageLabel(tk.Label):

    def load(self, im):
        if isinstance(im, str):
            im = Image.open(im)
        self.loc = 0
        self.frames = []

        try:
            for i in count(1):
                self.frames.append(ImageTk.PhotoImage(im.copy()))
                im.seek(i)
        except EOFError:
            pass

        try:
            self.delay = im.info['duration']
        except:
            self.delay = 100

        if len(self.frames) == 1:
            self.config(image=self.frames[0])
        else:
            self.next_frame()

    def unload(self):
        self.config(image=None)
        self.frames = None

    def next_frame(self):
        if self.frames:
            self.loc += 1
            self.loc %= len(self.frames)
            self.config(image=self.frames[self.loc])
            self.after(self.delay, self.next_frame)


def wit_vit():

    root1 = Toplevel()
    root1.state('zoomed')
    root1.title('vitness/informant')
    root1.configure(bg='#ffeb99')
    Fa1 = tk.Frame(root1, borderwidth=40, bg="black", relief=tk.SUNKEN)
    lb1 = ImageLabel(root1)
    lb1.pack(side=tk.TOP)
    lb1.load('Y_giff.gif')
    #pygame.init()
    pygame.mixer.music.load('a3.mp3')
    pygame.mixer.music.play()
    # time.sleep(5)
    # playsound('a3.mp3')
    # pygame.mixer.music.load('a3.mp3')
    # pygame.mixer.music.play()
    Fa1.pack()
    Fa2 = tk.Frame(root1, borderwidth=40, bg="#ffeb99", relief=tk.SUNKEN)
    lb2 = tk.Label(root1, text="YOU ARE",bg="#ffeb99")
    lb2.pack(side=tk.TOP)

    s = "welcome to virtual police station ."
    # file = "file1.mp3"
    # tts = gTTS(text=s, slow=False, lang='en')
    # tts.save(file)
    # os.system("start file1.mp3")

    but_witness= tk.Button(Fa2, text="Informant", width='63', height='5',command=per_wit)
    but_witness.pack(side=tk.LEFT, padx=150, fill=tk.X)
    but_victim= tk.Button(Fa2, text="Victim", width='63', height='5',command=per_vic)
    but_victim.pack(side=tk.RIGHT, padx=20, fill=tk.X)
    Fa2.pack(side=tk.BOTTOM, fill=tk.BOTH)

    root1.mainloop()

import speech_recognition as sr
#import pyaudio
from tkinter import messagebox
def funda():
    def takecommand():
        r = sr.Recognizer()
        with sr.Microphone() as source:
            print("Listening...")
            r.pause_threshold = 1
            audio = r.listen(source)
        try:
            print("Recognizing...")
            query = r.recognize_google(audio, language='en-in')
        except Exception as e:
            print("say that again")
            return "None"
        return query


def per_wit():
    root4 = Toplevel()
    root4.state('zoomed')
    root4.title('grid')
    root4.configure(bg='#ffeb99')


    label3_0 = tk.Label(root4, text="Personal details of witness", width=20,bg='#ffeb99' ,font=("bold", 23)).grid(row=0, column=0)
    lb313 = ImageLabel(root4)
    lb313.grid(row=0, column=4)
    lb313.load('Y_giff.gif')


    label3_1 = tk.Label(root4, text="Name", width=40, bg='#ffeb99', font=("bold", 13),padx=20,pady=20).grid(row=1, column=0)

    entry3_1 = tk.Entry(root4)
    entry3_1.grid(row=1, column=1)

    def play1_wt():
        pygame.mixer.music.load("a5.mp3")
        pygame.mixer_music.play()

    def rec1_wt():
        y1=funda()


    but3_q1 = tk.Button(root4, text="question1", width='10', height='1',command=play1_wt).grid(row=1, column=2)
    but3_r1 = tk.Button(root4, text="record1", width='10', height='1',command=rec1_wt).grid(row=1, column=3)


   # pygame.init()
    pygame.mixer.music.load('a4.mp3')
    pygame.mixer.music.play()
    label3_2 = tk.Label(root4, text="Contact Number", bg='#ffeb99', width=20, font=("bold", 13)).grid(row=2,column=0)


    entry3_2 = tk.Entry(root4)
    entry3_2.grid(row=2, column=1)
    def play2_wt():
        pygame.mixer.music.load('a10.mp3')
        pygame.mixer_music.play()
    def rec2_wt():
        y2=funda()



    but3_q2 = tk.Button(root4, text="question2", width='10', height='1',command=play2_wt).grid(row=2, column=2)
    but3_r2 = tk.Button(root4, text="record2", width='10', height='1',command=rec2_wt()).grid(row=2,column=3)

    label3_3 = tk.Label(root4, text="Gender", bg='#ffeb99', width=8, font=("bold", 13)).grid(row=3, column=0)
    y3 = tk.IntVar()
    tk.Radiobutton(root4, text="Male", padx=5, variable=y3, value=1, anchor=W, bg='#ffeb99', font=("bold", 12)).grid(
        row=3, column=1)
    tk.Radiobutton(root4, text="Female", padx=5, variable=y3, value=2, anchor=W, bg='#ffeb99', font=("bold", 12)).grid(
        row=3, column=2)
    tk.Radiobutton(root4, text="other", padx=20, variable=y3, value=3, anchor=W, bg='#ffeb99', font=("bold", 12)).grid(
        row=3, column=3)
    label3_4 = tk.Label(root4, text="Address", bg='#ffeb99', width=8, font=("bold", 13)).grid(row=4, column=0)

    entry3_4 = tk.Entry(root4)
    entry3_4.grid(row=4, column=1)
    def play4_wt():
        pygame.mixer.music.load('a9.mp3')
        pygame.mixer_music.play()
    def rec3_wt():
        y3=funda()

    but3_q4 = tk.Button(root4, text="question4", width='10', height='1',command=play4_wt).grid(row=4, column=2)
    but3_r4 = tk.Button(root4, text="record4", width='10', height='1',command=rec3_wt()).grid(row=4, column=3)

    label3_5 = tk.Label(root4, text=" Briefly explain the incidence", bg='#ffeb99', width=25, font=("bold", 13)).grid(row=5, column=0)

    entry3_5 = tk.Entry(root4)
    entry3_4.grid(row=5, column=1)
    def play5_wt():
        pygame.mixer.music.load('a14.mp3')
        pygame.mixer_music.play()
    def rec4_wt():
        y4=funda()

    but3_q5 = tk.Button(root4, text="question5", width='10', height='1',command=play5_wt).grid(row=5, column=2)
    but3_r5 = tk.Button(root4, text="record5", width='10', height='1',command=rec4_wt()).grid(row=5, column=3)

    label3_6 = tk.Label(root4, text="Description of accused", bg='#ffeb99', width=25, font=("bold", 13)).grid(row=6, column=0)

    entry3_6 = tk.Entry(root4)
    entry3_6.grid(row=6, column=1)
    def play6_wt():
        pygame.mixer.music.load("a20.mp3")
        pygame.mixer_music.play()

    but3_q6 = tk.Button(root4, text="question6", width='10', height='1',command=play6_wt).grid(row=6, column=2)
    but3_r6 = tk.Button(root4, text="record6", width='10', height='1').grid(row=6, column=3)

    label3_7 = tk.Label(root4, text="date of birth", bg='#ffeb99', width=20, font=("bold", 13)).grid(row=7, column=0)

    entry3_7 = tk.Entry(root4)
    entry3_7 .grid(row=7, column=1)

    but3_q7 = tk.Button(root4, text="question7", width='10', height='1').grid(row=7, column=2)
    but3_r7 = tk.Button(root4, text="record7", width='10', height='1').grid(row=7, column=3)

    # label_8 = tk.Label(root2, text=" Adhar Number", bg='#00cc99', width=8, font=("bold", 13)).grid(row=8, column=0)
    # x12 = StringVar()
    # entry_8 = tk.Entry(root2, textvariable=x12).grid(row=8, column=1)
    # but_q8 = tk.Button(root2, text="question8", width='10', height='1').grid(row=8, column=2)
    # but_r8 = tk.Button(root2, text="record8", width='10', height='1').grid(row=8, column=3)

    but3_submit = tk.Button(root4, text="Back", width='12', height='3', padx=5).grid(row=8, column=0)
    # but3_submit = tk.Button(root4, text="Add witness", width='12', height='3', padx=5).grid(row=8, column=1)

    def sub_m():
        y1 = entry3_1.get()
        print(y1)
        y2 = entry3_2.get()
        print(y2)
        y4 = entry3_4.get()
        print(y4)
        y5 = entry3_5.get()
        print(y5)
        y6 = entry3_6.get()
        print(y6)
        y7 = entry3_7.get()
        print(y7)


    def u():
        sub_m()
        root4.destroy()
        signature_wit()

    but3_submit = tk.Button(root4, text="submit and next", width='12', height='3', padx=5, command=u).grid(row=8, column=1)


    root4.mainloop()


def inci_vic():
    root3 = Toplevel()
    root3.state('zoomed')
    root3.title('grid')
    root3.configure(bg='#ffeb99')
    label2_0 = tk.Label(root3, text="DETAILS OF INCIDENCE ", width=20,bg='#ffeb99',font=("bold", 23),pady=15,padx=20).grid(row=0, column=0)
    lb213 = ImageLabel(root3)
    lb213.grid(row=0, column=4)
    lb213.load('Y_giff.gif')
    # pygame.init()
    pygame.mixer.music.load('a14.mp3')
    pygame.mixer.music.play()

    label2_1 = tk.Label(root3, text="Place of occurence", width=20, bg='#ffeb99', font=("bold", 13),pady=15,padx=20).grid(row=1, column=0)

    entry2_1 = tk.Entry(root3)
    entry2_1.grid(row=1, column=1)
    def play1_vinci():
        pygame.mixer.music.load("a15.mp3")
        pygame.mixer_music.play()

    but2_q1 = tk.Button(root3, text="Listen", width='10', height='1',command=play1_vinci).grid(row=1, column=2)
    but2_r1 = tk.Button(root3, text="Record", width='10', height='1').grid(row=1, column=3)

    pygame.mixer.music.load('a14.mp3')
    pygame.mixer.music.play()
    label2_2 = tk.Label(root3, text="Date of incidence", bg='#ffeb99', width=20, font=("bold", 13),pady=15,padx=20).grid(row=2,column=0)

    entry2_2 = tk.Entry(root3)
    entry2_2.grid(row=2, column=1)
    def play2_vinci():
        pygame.mixer.music.load("a16.mp3")
        pygame.mixer_music.play()
    but2_q2 = tk.Button(root3, text="Listen", width='10', height='1',command=play2_vinci).grid(row=2, column=2)
    but2_r2 = tk.Button(root3, text="Record", width='10', height='1').grid(row=2,column=3)

    label_3 = tk.Label(root3, text="", bg='#ffeb99', width=20, font=("bold", 13),pady=15,padx=20).grid(row=3, column=0)

    entry22_2 = tk.Entry(root3)
    entry22_2.grid(row=2, column=1)
    def play3_vinci():
        pygame.mixer.music.load("a17.mp3")
        pygame.mixer_music.play()
    but22_q2 = tk.Button(root3, text="Listen", width='10', height='1',command=play3_vinci).grid(row=2, column=2)
    but22_r2 = tk.Button(root3, text="Record", width='10', height='1').grid(row=2, column=3)
    # var = tk.IntVar()
    # tk.Radiobutton(root2, text="Male", padx=5, variable=var, value=1, anchor=W, bg='#00cc99', font=("bold", 12)).grid(
    #     row=3, column=1)
    # tk.Radiobutton(root2, text="Female", padx=5, variable=var, value=2, anchor=W, bg='#00cc99', font=("bold", 12)).grid(
    #     row=3, column=2)
    # tk.Radiobutton(root2, text="other", padx=20, variable=var, value=3, anchor=W, bg='#00cc99', font=("bold", 12)).grid(
    #     row=3, column=3)
    label2_4 = tk.Label(root3, text="Nature of offence", bg='#ffeb99', width=20, font=("bold", 13),pady=15,padx=20).grid(row=4, column=0)

    entry2_4 = tk.Entry(root3)
    entry2_4.grid(row=4, column=1)
    def play4_vinci():
        pygame.mixer.music.load("a18.mp3")
        pygame.mixer_music.play()
    but2_q4 = tk.Button(root3, text="Listen", width='10', height='1',command=play4_vinci).grid(row=4, column=2)
    but2_r4 = tk.Button(root3, text="Record", width='10', height='1').grid(row=4, column=3)

    label2_5 = tk.Label(root3, text=" Particulars of property", bg='#ffeb99', width=20, font=("bold", 13),pady=15,padx=20).grid(row=5, column=0)

    entry2_5 = tk.Entry(root3)
    entry2_5.grid(row=5, column=1)
    def play5_vinci():
        pygame.mixer.music.load("a19.mp3")
        pygame.mixer_music.play()
    but2_q5 = tk.Button(root3, text="Listen", width='10', height='1',command=play5_vinci).grid(row=5, column=2)
    but2_r5 = tk.Button(root3, text="Record", width='10', height='1').grid(row=5, column=3)

    label2_6 = tk.Label(root3, text="Description of accused", bg='#ffeb99', width=20, font=("bold", 13),pady=15,padx=20).grid(row=6, column=0)

    entry2_6 = tk.Entry(root3)

    entry2_6 .grid(row=6, column=1)
    def play6_vinci():
        pygame.mixer.music.load("a20.mp3")
        pygame.mixer_music.play()
    but2_q6 = tk.Button(root3, text="Listen", width='10', height='1',command=play6_vinci).grid(row=6, column=2)
    but2_r6 = tk.Button(root3, text="Record", width='10', height='1').grid(row=6, column=3)

    label2_7 = tk.Label(root3, text=" Complain", bg='#ffeb99', width=20, font=("bold", 13),pady=15,padx=20).grid(row=7, column=0)

    entry2_7 = tk.Entry(root3)
    entry2_7 .grid(row=7, column=1)
    def play7_vinci():
        pygame.mixer.music.load("a22.mp3")
        pygame.mixer_music.play()
    but2_q7 = tk.Button(root3, text="Listen", width='10', height='1',command=play7_vinci).grid(row=7, column=2)
    but2_r7 = tk.Button(root3, text="Record", width='10', height='1').grid(row=7, column=3)


    but2_submit = tk.Button(root3, text="Back", width='12', height='3', padx=1,command=root3.destroy).grid(row=8, column=0)

    def sub_mv_inci():
        global x15,x16,x17,x18,x19,x20,x22
        x15 = entry2_1.get()
        print(x15)
        x16 = entry2_2.get()
        print(x16)
        x17 = entry22_2.get()
        print(x17)
        x18 = entry2_4.get()
        print(x18)
        x19 = entry2_5.get()
        print(x19)
        x20 = entry2_6.get()
        print(x20)
        x22 = entry2_7.get()
        print(x22)


    def u1():
        sub_mv_inci()
        root3.destroy()
        signature_victim()


    but2_submit = tk.Button(root3, text="submit and next", width='12', height='3', padx=1, command=u1).grid(row=8,column=2)
    root3.mainloop()

def per_vic():
    root2 = Toplevel()
    root2.state('zoomed')
    root2.title('personal details of victim')
    root2.configure(bg='#ffeb99')
    pygame.mixer.music.load('a4.mp3')
    pygame.mixer.music.play()

    label_0 = tk.Label(root2, text="PERSONAL DETAILS ", width=20,bg='#ffeb99', font=("bold", 23)).grid(row=0, column=0)
    lb13 = ImageLabel(root2)
    lb13.grid(row=0, column=4)
    lb13.load('Y_giff.gif')

    label_1 = tk.Label(root2, text="FullName", width=20, bg='#ffeb99', font=("bold", 16),pady=15).grid(row=1, column=0)

    entry_1 = tk.Entry(root2)
    entry_1.grid(row=1, column=1)
    def play1_perv():
        pygame.mixer.music.load("a5.mp3")
        pygame.mixer_music.play()
    but_q1 = tk.Button(root2, text="Listen", width='10', height='1',command=play1_perv).grid(row=1, column=2)
    but_r1 = tk.Button(root2, text="Record", width='10', height='1').grid(row=1, column=3)


    label_2 = tk.Label(root2, text="Father's/Mother's name", bg='#ffeb99', width=20, font=("bold", 13),pady=15).grid(row=2,column=0)

    entry_2 = tk.Entry(root2)
    entry_2.grid(row=2, column=1)
    def play2_perv():
        pygame.mixer.music.load("a6.mp3")
        pygame.mixer_music.play()
    but_q2 = tk.Button(root2, text="Listen", width='10', height='1',command=play2_perv).grid(row=2, column=2)
    but_r2 = tk.Button(root2, text="Record", width='10', height='1').grid(row=2,column=3)
    label_3 = tk.Label(root2, text="Gender", bg='#ffeb99', width=8, font=("bold", 16),pady=15).grid(row=3, column=0)
    var = tk.IntVar()
    tk.Radiobutton(root2, text="Male", padx=5, variable=var, value=1, anchor=W, bg='#ffeb99', font=("bold", 12)).grid(
        row=3, column=1)
    tk.Radiobutton(root2, text="Female", padx=5, variable=var, value=2, anchor=W, bg='#ffeb99', font=("bold", 12)).grid(
        row=3, column=2)
    tk.Radiobutton(root2, text="other", padx=20, variable=var, value=3, anchor=W, bg='#ffeb99', font=("bold", 12)).grid(
        row=3, column=3)
    label_4 = tk.Label(root2, text="Age:", bg='#ffeb99', width=20, font=("bold", 13),pady=15).grid(row=4, column=0)

    entry_4 = tk.Entry(root2)
    entry_4 .grid(row=4, column=1)

    def play4_perv():
        pygame.mixer.music.load("a8.mp3")
        pygame.mixer_music.play()
    but_q4 = tk.Button(root2, text="Listen", width='10', height='1',command= play4_perv).grid(row=4, column=2)
    but_r4 = tk.Button(root2, text="Record", width='10', height='1').grid(row=4, column=3)

    label_5 = tk.Label(root2, text=" Address", bg='#ffeb99', width=20, font=("bold", 13)).grid(row=5, column=0)

    entry_5 = tk.Entry(root2)
    entry_5.grid(row=5, column=1)
    def play5_perv():
        pygame.mixer.music.load("a9.mp3")
        pygame.mixer_music.play()
    but_q5 = tk.Button(root2, text="Listen", width='10', height='1',command=play5_perv).grid(row=5, column=2)
    but_r5 = tk.Button(root2, text="Record", width='10', height='1').grid(row=5, column=3)

    label_6 = tk.Label(root2, text="Contact number", bg='#ffeb99', width=20, font=("bold", 13),pady=15).grid(row=6, column=0)

    entry_6 = tk.Entry(root2)
    entry_6.grid(row=6, column=1)
    def play6_perv():
        pygame.mixer.music.load("a10.mp3")
        pygame.mixer_music.play()
    but_q6 = tk.Button(root2, text="Listen", width='10', height='1',command=play6_perv).grid(row=6, column=2)
    but_r6 = tk.Button(root2, text="Record", width='10', height='1').grid(row=6, column=3)

    label_7 = tk.Label(root2, text=" Email", bg='#ffeb99', width=20, font=("bold", 13)).grid(row=7, column=0)

    entry_7 = tk.Entry(root2)
    entry_7.grid(row=7, column=1)
    def play7_perv():
        pygame.mixer.music.load("a11.mp3")
        pygame.mixer_music.play()
    but_q7 = tk.Button(root2, text="Listen", width='10', height='1',command=play7_perv).grid(row=7, column=2)
    but_r7 = tk.Button(root2, text="Record", width='10', height='1').grid(row=7, column=3)

    label_8 = tk.Label(root2, text=" date of birth", bg='#ffeb99', width=20, font=("bold", 13),pady=15).grid(row=8, column=0)

    entry_8 = tk.Entry(root2)
    entry_8.grid(row=8, column=1)
    # def play8_perv():
    #     pygame.mixer.music.load("a11.mp3")
    #     pygame.mixer_music.play()
    but_q8 = tk.Button(root2, text="Listen", width='10', height='1').grid(row=8, column=2)
    but_r8 = tk.Button(root2, text="Record", width='10', height='1').grid(row=8, column=3)

    but_submit = tk.Button(root2, text="BACK", width='12', height='3', padx=5,command=wit_vit).grid(row=9, column=0)
    but_submit = tk.Button(root2, text="SUBMIT & NEXT", width='12', height='3', padx=5,command=inci_vic).grid(row=9, column=1)

    def sub_pervict():
        global x5,x6,x8,x9,x10,x11,x12
        x5 = entry_1.get()
        print(x5)
        x6 = entry_2.get()
        print(x6)
        # x17 = entry22_2.get()
        # print(x17)
        x8 = entry_4.get()
        print(x8)
        x9 = entry_5.get()
        print(x9)
        x10 = entry_6.get()
        print(x10)
        x11 = entry_7.get()
        print(x11)
        x12 = entry_8
        print(x12)
    def u2():
        sub_pervict()
        inci_vic()

    but_submit = tk.Button(root2, text="submit and next", width='12', height='3', padx=5, command=u2).grid(row=9,column=1)

    root2.mainloop()


def adhar():

    root_a=Toplevel()
    root_a.state('zoomed')
    root_a.title('Aadhaar details')
    root_a.configure(bg='#ffeb99')
    pygame.mixer.music.load('a12.mp3')
    pygame.mixer.music.play()
    def document():
        doc = docx.Document()
        doc.add_paragraph(x5 + ', ' + x6)
        doc.add_paragraph(x9)
        doc.add_paragraph(x10)
        doc.add_paragraph(x11)
        doc.add_paragraph(x16)
        doc.add_paragraph()
        doc.add_paragraph("To,")
        doc.add_paragraph("The Police Officer In charge")
        doc.add_paragraph("Ranjhi Police thana")
        doc.add_paragraph()
        doc.add_paragraph("Sub: " +  x22)
        doc.add_paragraph("Respected Sir,")
        doc.add_paragraph()
        doc.add_paragraph("I would like to bring the following facts to your kind notice:")
        doc.add_paragraph(x20 +x22)
        doc.add_paragraph("..............................................................")
        doc.add_paragraph("I seek your help and request to kindly register my F.I.R. in the subject matter.")
        doc.add_paragraph()
        doc.add_paragraph("For your Ready Reference I verified aadhar number with online verification")
        doc.add_paragraph()
        doc.add_paragraph("Hope You will do the needful favourably at the earliest.")
        doc.add_paragraph()
        doc.add_paragraph("Thanks and regards")
        doc.add_paragraph("Yours Sincerely")
        doc.add_picture('victim_sign.png')
        doc.save('test.docx')
        os.system('start test.docx')
    labela_0 = tk.Label(root_a, text="AADHAR DETAILS ", width=20,bg='#ffeb99', font=("bold", 23)).grid(row=0, column=0)
    lba13 = ImageLabel(root_a)
    lba13.grid(row=0, column=4)
    lba13.load('Y_giff.gif')
    labela_1 = tk.Label(root_a, text="Adhar Number", width=20, bg='#ffeb99', font=("bold", 13)).grid(row=1, column=0)
    entrya_1 = tk.Entry(root_a)
    entrya_1.grid(row=1, column=1)
    buta_q1 = tk.Button(root_a, text="Listen", width='10', height='1').grid(row=1, column=2)
    buta_r1 = tk.Button(root_a, text="Record", width='10', height='1').grid(row=1, column=3)
    labela_2 = tk.Label(root_a, text="Linked Mobile Number", bg='#ffeb99', width=20, font=("bold", 13)).grid(row=2, column=0)
    entrya_2 = tk.Entry(root_a)
    entrya_2.grid(row=2, column=1)
    def verify():
        x1 = entrya_1.get()
        x1 = str(x1)
        y2 = entrya_2.get()
        y2 = str(y2)
        x = auth(aadhar_number=x1,mobile_no=y2)
        def document():
            doc = docx.Document()
            doc.add_paragraph(x5 + ', ' + x6)
            doc.add_paragraph(x9)
            doc.add_paragraph(x10)
            doc.add_paragraph(x11)
            doc.add_paragraph(x16)
            doc.add_paragraph()
            doc.add_paragraph("To,")
            doc.add_paragraph("The Police Officer In charge")
            doc.add_paragraph("Ranjhi Police thana")
            doc.add_paragraph()
            doc.add_paragraph("Sub: " +  x22)
            doc.add_paragraph("Respected Sir,")
            doc.add_paragraph()
            doc.add_paragraph("I would like to bring the following facts to your kind notice:")
            doc.add_paragraph(x20 +x22)
            doc.add_paragraph("..............................................................")
            doc.add_paragraph("I seek your help and request to kindly register my F.I.R. in the subject matter.")
            doc.add_paragraph()
            doc.add_paragraph("For your Ready Reference I verified aadhar number with online verification")
            doc.add_paragraph()
            doc.add_paragraph("Hope You will do the needful favourably at the earliest.")
            doc.add_paragraph()
            doc.add_paragraph("Thanks and regards")
            doc.add_paragraph("Yours Sincerely")
            doc.add_picture('victim_sign.png')
            doc.save('test.docx')
            os.system('start test.docx')
        if x :
            document()
        else :
            messagebox.showinfo("Not Verified")
        #messagebox.showinfo("not verified")
    buta_q2 = tk.Button(root_a, text="Listen", width='10', height='1').grid(row=2, column=2)
    buta_r2 = tk.Button(root_a, text="Record", width='10', height='1').grid(row=2, column=3)
    click = tk.Button(root_a,text ="Submit", width = '20',height ='1',command=verify).grid(row=3,column=2)

    def auth(aadhar_number,mobile_no):
        driver = webdriver.Chrome("chromedriver")
        url = "https://resident.uidai.gov.in/verify-email-mobile"
        driver.get(url)
        driver.find_element_by_id('verify_uidno').send_keys(aadhar_number)
        driver.find_element_by_id ('verify_mobile').send_keys(mobile_no)
        driver.find_element_by_id('smt_btn').click()
        time.sleep(40)
        myScreenshot = pyautogui.screenshot()

        myScreenshot.save(r'screenshot1.png')
        driver.quit()
        path = 'screenshot1.png'
        img = Image.open(path)
        img = img.convert('RGBA')
        pix = img.load()
        for y in range(img.size[1]):
            for x in range(img.size[0]):
                if pix[x, y][0] < 102 or pix[x, y][1] < 102 or pix[x, y][2] < 102:
                    pix[x, y] = (0, 0, 0, 255)
                else:
                    pix[x, y] = (255, 255, 255, 255)
        img.save('temp.png')
        im = Image.open(r"temp.png")
        left = 50
        top = 450
        right = 700
        bottom = 550
        im1 = im.crop((left, top, right, bottom))
        pytesseract.pytesseract.tesseract_cmd = "C:\Program Files (x86)\Tesseract-OCR\\tesseract.exe"
        text = pytesseract.image_to_string(im1)
        #os.remove('temp.jpg')
        if 'matches with our' or 'verified with our' in text :
            return True
        else :
            im = Image.open(r"temp.png")
            left = 100
            top = 500
            right = 900
            bottom = 550
            im1 = im.crop((left, top, right, bottom))
            pytesseract.pytesseract.tesseract_cmd = "C:\Program Files (x86)\Tesseract-OCR\\tesseract.exe"
            text = pytesseract.image_to_string(im1)
            if 'matches with our' or 'verified with our' in text :
                return True
            else :
                return False

root = Tk()
root.state('zoomed')
root.title('Virtual Police Station')

F1 = Frame(root, borderwidth=7, bg="#ffeb99", relief=SUNKEN)
button_english = Button(F1, text="English", width='9', height='10', command = wit_vit)
button_english.pack(side=TOP, pady=9, fill=Y)
button_hindi = Button(F1, text="Hindi", width='9', height='10',command=Hwit_vit)
pygame.init()

def signature_victim():
    def save():
        filename = 'victim_sign.png'   # image_number increments by 1 at every save
        image1.save(filename)
    def activate_paint(e):
        global lastx, lasty
        cv.bind('<B1-Motion>', paint)
        lastx, lasty = e.x, e.y
    def paint(e):
        global lastx, lasty
        x, y = e.x, e.y
        cv.create_line((lastx, lasty, x, y), width=1)
        draw.line((lastx, lasty, x, y), fill='black', width=1)
        lastx, lasty = x, y
    sign = Tk()
    lastx, lasty = None, None
    image_number = 0
    cv = Canvas(sign, width=440, height=380, bg='white')
# --- PIL
    image1 = Image.new('RGB', (440, 380), 'white')
    draw = ImageDraw.Draw(image1)

    cv.bind('<1>', activate_paint)
    cv.pack(expand=YES, fill=BOTH)

    def lolo():
        save()
        sign.destroy()
        adhar()
    sign.title("Please Give Your Signature")
    btn_save = Button(sign, text="Save & Next", command=lolo)
    btn_save.pack()
    sign.mainloop()
def signature_wit():
    def lolo():
        save()
        sign.destroy()
        adhar()
    def save():
        filename = 'Witness_sign.png'   # image_number increments by 1 at every save
        image1.save(filename)
    def activate_paint(e):
        global lastx, lasty
        cv.bind('<B1-Motion>', paint)
        lastx, lasty = e.x, e.y
    def paint(e):
        global lastx, lasty
        x, y = e.x, e.y
        cv.create_line((lastx, lasty, x, y), width=1)
        draw.line((lastx, lasty, x, y), fill='black', width=1)
        lastx, lasty = x, y
    sign = Tk()
    lastx, lasty = None, None
    image_number = 0
    cv = Canvas(sign, width=240, height=180, bg='white')
    image1 = Image.new('RGB', (240, 180), 'white')
    draw = ImageDraw.Draw(image1)

    cv.bind('<1>', activate_paint)
    cv.pack(expand=YES, fill=BOTH)
    sign.title("Please Give Your Signature")
    btn_save = Button(sign, text="Save & Next", command=lolo)
    btn_save.pack()
    sign.mainloop()

pygame.mixer.music .load('a1.mp3')
pygame.mixer.music.play()

button_hindi.pack(side=BOTTOM, pady=9, fill=Y)
F1.pack(side=LEFT, anchor=W, fill=BOTH)

F2 = Frame(root, borderwidth=40, bg="blue", relief=SUNKEN)
image = Image.open('backimage.png')
photo = ImageTk.PhotoImage(image)
label1 = Label(image=photo, height=1000, width=1000)
label1.pack(fill=BOTH)
F2.pack(side=TOP, fill=Y)
root.mainloop()
